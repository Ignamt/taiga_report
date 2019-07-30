"""Contains classes for the different report printers"""
import datetime as dt
from pathlib import Path
from docx import Document


class Printer:
    """Base class for printers.

    This class has the _check_filename() method and should contain
    any methods that will be used by all printer classes.

    THIS IS AN ABSTRACT CLASS AND SHOULD NOT BE USED ON IT'S OWN

    """

    ext = None

    @classmethod
    def _check_filename(cls, project):
        """Generate new report filename.

        Checks the file system if there is a file for the current month's
        report. If there is one, makes a new version. If there is none,
        generates the filename and returns it.

        ARGS:
            - project: str() of the project title.

        RETURNS: filename

        """
        if not cls.ext:
            raise Exception("No report extension was detected. Use a specific"
                            " format printer instead")
        year = dt.date.today().year
        month = str(dt.date.today().month).rjust(2, "0")
        filename = project + "_report_{}-{}".format(month, year)
        if Path(filename + cls.ext).is_file():
            filename += "_1"
        while Path(filename + cls.ext).is_file():
            filename = filename.replace(filename[-1], str(int(filename[-1])+1))
        filename += cls.ext
        return filename


class MarkdownPrinter(Printer):
    """Prints the report in markdown format.

    To use:
        MarkdownPrinter.print_markdown(report)

    """

    ext = ".md"

    @classmethod
    def print_markdown(cls, report):
        """Print the report in markdown format into a file.

        ARGS:
            - report: Report() object containing the US info

        """
        filename = cls._check_filename(report.project)
        with open(filename, "a+") as file:
            file.write(cls.md_title(report.project))
            for section in report._report_sections:
                if section in report._report:
                    cls._print_section_md(section, file, report)

    @classmethod
    def _print_section_md(cls, section, file, report):
        """Write a section with it's epics and US to a file.

        PARAMETERS:
            - section: string of a section of the report.
            - file: file-like object to which to print to.
            - report: Report() object containing the US info

        """
        file.write(cls.md_section(section))
        rep_section = report._report[section]
        if rep_section.get("user_stories"):
            cls._print_userstories_md(rep_section["user_stories"], file)

        for epic in report._report[section]:
            if epic == "user_stories":
                continue

            cls._print_epic_md(section, epic, file, report)

    @classmethod
    def _print_epic_md(cls, section, epic, file, report):
        """Write an epic with it's US to a section in a file.

        PARAMETERS:
            - section: String of a section of the report
            - epic: String of an epic that goes in the received section
            - file: file-like object to which to print to
            - report: Report() object containing the US info

        """
        file.write(cls.md_epic(epic))
        cls._print_userstories_md(report._report[section][epic], file)

    @classmethod
    def _print_userstories_md(cls, userstories, file):
        """Write US to a file.

        PARAMETERS:
            - userstories: List of US
            - file: file-like object to which to print to

        """
        for us in userstories:
            file.write(cls.md_user_story(us))
        # Add one final newline to separate from other parts of the report
        file.write("\n")

    @classmethod
    def md_title(cls, content):
        """Format content string as Markdown h1.

        PARAMETERS:
            - content: string with the project title.

        RETURNS: str formatted as a <h1> for the title.

        """
        return "# {}\n\n".format(content)

    @classmethod
    def md_section(cls, content):
        """Format content string as Markdown h2.

        PARAMETERS:
            - content: string with the project section.

        RETURNS: str formatted as a <h2> for the section.

        """
        return "## {}\n\n".format(content.capitalize())

    @classmethod
    def md_epic(cls, content):
        """Format content string as Markdown h3.

        PARAMETERS:
            - content: string with the project epic.

        RETURNS: str formatted as a <h3> for the epic.

        """
        return "### {}\n\n".format(content.capitalize())

    @classmethod
    def md_user_story(cls, content):
        """Format content string as Markdown list item.

        PARAMETERS:
            - content: string with the project user story.

        RETURNS: str formatted as a  list for the user story.

        """
        return "* {}\n".format(content.capitalize())


class DocxPrinter(Printer):
    """Prints the report in docx format.

    To use:
        DocxPrinter.print_markdown(report)

    """

    ext = ".docx"

    @classmethod
    def print_docx(cls, report):
        """Print the report in Microsoft Word format into a file.

        ARGS:
            - report: Report() object containing the US info

        """
        filename = cls._check_filename(report.project)
        document = Document()
        cls.docx_title(document, report.project)
        for section in report._report_sections:
            if section in report._report:
                cls._print_section_docx(section, document, report)
        document.save(filename)

    @classmethod
    def _print_section_docx(cls, section, document, report):
        """Write a section with it's epics and US to a Document().

        PARAMETERS:
            - section: string of a section of the report.
            - document: docx.Document() object to which the function prints to.
            - report: Report() object containing the US info.

        """
        cls.docx_section(document, section)
        rep_section = report._report[section]
        if rep_section.get("user_stories"):
            cls._print_userstories_docx(rep_section["user_stories"], document)

        for epic in report._report[section]:
            if epic == "user_stories":
                continue

            cls._print_epic_docx(section, epic, document, report)

    @classmethod
    def _print_epic_docx(cls, section, epic, document, report):
        """Write an epic with it's US to a section in a Document().

        PARAMETERS:
            - section: String of a section of the report.
            - epic: String of an epic that goes in the received section.
            - document: docx.Document() object to which the function prints to.
            - report: Report() object containing the US info.

        """
        cls.docx_epic(document, epic)
        cls._print_userstories_docx(report._report[section][epic], document)

    @classmethod
    def _print_userstories_docx(cls, userstories, document):
        """Write US to a Document().

        PARAMETERS:
            - userstories: List of US
            - document: docx.Document() object to which the function prints to.

        """
        for us in userstories:
            cls.docx_user_story(document, us)

    @classmethod
    def docx_title(cls, document, content):
        """Format content string as Microsoft Word Title.

        PARAMETERS:
            - document: docx.Document() object to which the function prints to.
            - content: string with the project title.

        RETURNS: str formatted as a Title for the title.

        """
        return document.add_heading(content.capitalize(), level=0)

    @classmethod
    def docx_section(cls, document, content):
        """Format content string as Microsoft Word Heading 1.

        PARAMETERS:
            - document: docx.Document() object to which the function prints to.
            - content: string with the project section.

        RETURNS: str formatted as a Heading 1 for the section.

        """
        return document.add_heading(content.capitalize(), level=1)

    @classmethod
    def docx_user_story(cls, document, content):
        """Format content string as Microsoft Word Bullet List Item.

        PARAMETERS:
            - document: docx.Document() object to which the function prints to.
            - content: string with the project epic.

        RETURNS: str formatted as a Heading 2 for the epic.

        """
        return document.add_paragraph(content.capitalize(),
                                      style='List Bullet 2')

    @classmethod
    def docx_epic(cls, document, content):
        """Format content string as Microsoft Word Heading 2.

        PARAMETERS:
            - document: docx.Document() object to which the function prints to.
            - content: string with the project user story.

        RETURNS: str formatted as a  list for the user story.

        """
        return document.add_heading(content.capitalize(), level=5)
