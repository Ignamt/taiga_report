"""Contains classes for the different sections of the report."""
import datetime as dt
from pathlib import Path
from docx import Document

class UserStory:
    """Contains all the US info needed for the report."""

    def __init__(self, us):
        """Set up attributes for the instance.

        Parameters:
        - us: Dict with all the info that comes from the Taiga
            API

        """
        self.subject = us["subject"]
        self.epic = us["epics"][0]["subject"] if us["epics"] else []
        self.tags = us["tags"][0] if us["tags"] else []
        self.section = self._section
        self.subtasks = us["tasks"]
        # if us["due_date"]:
        #     self.due_date = dt.datetime.strptime(us["due_date"],
        #                                          "%Y-%m-%dT%H:%M:%S.%fZ")
        # else:
        #     self.due_date = None

    @property
    def _section(self):
        for tag in self.tags:
            if tag in ["general", "expedientes", "remitos", "administracion"]:
                return tag
            else:
                return None


class Report:
    """Gathers all sections and prints itself in various formats."""

    def __init__(self, project, yaml_dict):
        """Set up attributes for the instance."""
        self._report = dict()
        self.project = project.upper()
        self._report_sections = yaml_dict[project]["report_sections"]

    def classify_user_story(self, us):
        """Classify a US and store it in the report json.

        PARAMETERS:
            - us: UserStory() object

        EXAMPLE OF REPORT JSON:
        report = {
            "section-1": {
                "user_stories": ["US-1", "US-2"],
                "epic-1": ["US-3", "US-4"]
                "epic-2": ["US-5", "US-6"]
            },
            "section-2": {
                "user_stories": ["US-7", "US-8"]
            }
        }

        """
        if us.section not in self._report:
            self._report[us.section] = {"user_stories": []}
            section = self._report[us.section]
            if us.epic:
                section[us.epic] = [us.subject]
            else:
                section["user_stories"].append(us.subject)
        else:
            if not us.epic:
                self._report[us.section]["user_stories"].append(us.subject)

            elif us.epic not in self._report[us.section]:
                self._report[us.section].update({us.epic: [us.subject]})

            else:
                self._report[us.section][us.epic].append(us.subject)
        
    def _check_filename(self, ext):
        """Generate new report filename.

        Checks the file system if there is a file for the current month's
        report. If there is one, makes a new version. If there is none,
        generates the filename and returns it.

        RETURNS: filename

        """

        year = dt.date.today().year
        month = str(dt.date.today().month).rjust(2, "0")
        filename = self.project + "_report_{}-{}".format(month, year)
        if Path(filename+ext).is_file():
            filename += "_1"
        while Path(filename+ext).is_file():
            filename = filename.replace(filename[-1], str(int(filename[-1])+1))
        filename += ext
        return filename


class MarkdownPrinter:
    """Prints the report in markdown format.

    Parameters:
        report: Report instance to print in markdown.

    To use:
        MarkdownPrinter.print_markdown(report)

    """

    @classmethod
    def print_markdown(cls, report):
        """Print the report in markdown format into a file."""
        filename = report._check_filename(".md")
        with open(filename, "a+") as file:
            file.write(cls.md_title(report.project))
            for section in report._report_sections:
                if section in report._report:
                    cls._print_section_md(section, file, report)

    @classmethod
    def _print_section_md(cls, section, file, report):
        """Write a section with it's epics and US to a file.

        PARAMETERS:
            - section: string of a section of the report.
            - file: file-like object to which to print to.

        """
        file.write(cls.md_section(section))
        rep_section = report._report[section]
        if rep_section.get("user_stories"):
            cls._print_userstories_md(rep_section["user_stories"], file)

        for epic in report._report[section]:
            if epic == "user_stories":
                continue

            cls._print_epic_md(section, epic, file, report)

    @classmethod
    def _print_epic_md(cls, section, epic, file, report):
        """Write an epic with it's US to a section in a file.

        PARAMETERS:
            - section: String of a section of the report
            - epic: String of an epic that goes in the received section
            - file: file-like object to which to print to

        """
        file.write(cls.md_epic(epic))
        cls._print_userstories_md(report._report[section][epic], file)

    @classmethod
    def _print_userstories_md(cls, userstories, file):
        """Write US to a file.

        PARAMETERS:
            - userstories: List of US

        """
        for us in userstories:
            file.write(cls.md_user_story(us))
        # Add one final newline to separate from other parts of the report
        file.write("\n")
    
    @classmethod
    def md_title(cls, content):
        return "# {}\n\n".format(content)
    
    @classmethod
    def md_section(cls, content):
        return "## {}\n\n".format(content.capitalize())

    @classmethod
    def md_epic(cls, content):
        return "### {}\n\n".format(content.capitalize())
    
    @classmethod
    def md_user_story(cls, content):
        return "* {}\n".format(content.capitalize())

class DocxPrinter:
    """Prints the report in docx format."""
    
    @classmethod
    def print_docx(cls, report):
        filename = report._check_filename(".docx")
        document = Document()
        cls.docx_title(document, report.project)
        for section in report._report_sections:
            if section in report._report:
                cls._print_section_docx(section, document, report)  
        document.save(filename)
    
    @classmethod
    def _print_section_docx(cls, section, document, report):
        cls.docx_section(document, section)
        rep_section = report._report[section]
        if rep_section.get("user_stories"):
            cls._print_userstories_docx(rep_section["user_stories"], document)

    @classmethod
    def md_title(cls, content):
        """Format content string as Markdown h1.

        PARAMETERS:
            - content: string with the project title.

        RETURNS: str formatted as a <h1> for the title.

        """
        return "# {}\n\n".format(content)

    @classmethod
    def md_section(cls, content):
        """Format content string as Markdown h2.

        PARAMETERS:
            - content: string with the project section.

        RETURNS: str formatted as a <h2> for the section.

        """
        return "## {}\n\n".format(content.capitalize())

    @classmethod
    def md_epic(cls, content):
        """Format content string as Markdown h3.

        PARAMETERS:
            - content: string with the project epic.

        RETURNS: str formatted as a <h3> for the epic.

        """
        return "### {}\n\n".format(content.capitalize())

    @classmethod
    def md_user_story(cls, content):
        """Format content string as Markdown list item.

        PARAMETERS:
            - content: string with the project user story.

        RETURNS: str formatted as a  list for the user story.

        """
        return "* {}\n".format(content.capitalize())


class DocxPrinter:
    """Prints the report in docx format.

    Parameters:
        report: Report instance to print in docx.

    To use:
        DocxPrinter.print_markdown(report)
    """

    @classmethod
    def print_docx(cls, report):
        filename = report._check_filename(".docx")
        document = Document()
        cls.docx_title(document, report.project)
        for section in report._report_sections:
            if section in report._report:
                cls._print_section_docx(section, document, report)
        document.save(filename)

    @classmethod
    def _print_section_docx(cls, section, document, report):
        cls.docx_section(document, section)
        rep_section = report._report[section]
        if rep_section.get("user_stories"):
            cls._print_userstories_docx(rep_section["user_stories"], document)

        for epic in report._report[section]:
            if epic == "user_stories":
                continue

            cls._print_epic_docx(section, epic, document, report)

    @classmethod
    def _print_epic_docx(cls, section, epic, document, report):
        cls.docx_epic(document, epic)
        cls._print_userstories_docx(report._report[section][epic], document)
    
    @classmethod
    def _print_userstories_docx(cls, userstories, document):
        for us in userstories:
            cls.docx_user_story(document, us)

    @classmethod
    def docx_title(cls, document, content):
        return document.add_heading(content.capitalize(), level=0)

    @classmethod
    def docx_section(cls, document, content):
        return document.add_heading(content.capitalize(), level=1)

    @classmethod
    def docx_user_story(cls, document, content):
        return document.add_paragraph(content.capitalize(), style='List Bullet 2')

    @classmethod
    def docx_epic(cls, document, content):
        return document.add_heading(content.capitalize(), level=5)
